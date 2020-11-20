class ParseDOCXSpecial:
    def __init__(self, docx_filename, settings):
        from docx import Document
        self.filename = docx_filename
        self.output_filename = docx_filename.split(".")[0] + "_out.docx"
        self.doc = Document(self.filename)
        self.settings = settings
        self.placeholder = "PLACEHOLDER-"
        self.placeholder_num = 0
        
    def getPlaceholder(self):
        placeholder = f"{self.placeholder}[{self.placeholder_num}]"
        self.placeholdernum += 1
        return placeholder

    def isTextWantedFromRun(run, font_color=None):
        if font_color is not None:
            if not isCorrectFontColor(run, font_color):
                return False
        return True

    def isCorrectFontColor(run, font_color):
        return str(run.font.color.rgb) == font_color
    
    def getText(self):
        return self.paragraphs_text + self.tables_text + self.shapes_text

    def processParagraphs(self, font_color=None):
        self.paragraphs_text = ""
        print("Processing Paragraphs...")
        for paragraph in self.doc.paragraphs:
            try:
                for run in paragraph.runs:
                    if isTextWantedFromRun(run, font_color):
                        self.paragraphs_text += run.text
                        run.text = self.getPlaceholder()

            except Exception as e:
                print(f"Error: {repr(e)}")

            finally:        
                self.paragraphs_text += "\n"

        print("Done!")

    def processTables(self, font_color=None):
        self.tables_text = ""
        print("Processing Tables...")
        for table in self.doc.tables:
            try:
                for cell in table._cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if isTextWantedFromRun(run, font_color):
                                self.tables_text += run.text
                                self.tables_text += "\n"
                                run.text = self.getPlaceholder()

            except Exception as e:
                print(f"Error: {repr(e)}")

            finally:
                self.tables_text += "\n"

        print("Done!")

    def processShapes(self, font_color=None):
        self.shapes_text = ""
        print("Processing Shapes...")
        for p in doc.inline_shapes:
            print(p.text)
        
        print("Done!")

    def processDocument(self):
        print(f"Processing file {self.filename}")
        self.processParagraphs(**self.settings)
        self.processTables(**self.settings)
        self.processShapes(**self.settings)
        print(self.getText())
        # text = paragraph_text + table_text + shapes_text
        print(f"Saving file to {self.output_filename}")
        self.doc.save(self.output_filename)
        print(f"Done!")
